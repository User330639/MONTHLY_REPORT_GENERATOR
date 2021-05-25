import PIL
from PIL import Image
from docx import Document
import os
from docx.shared import Cm

document = Document()
old_cwd = os.getcwd()

os.chdir('./report_src')
for dir in os.listdir('./'):
    if os.path.isdir(dir):
        os.chdir(dir)
        print('Scanning directory: ' + os.getcwd())
        text_file = open('Text.txt', 'r')
        text = text_file.read() + '\n'
        text_file.close()
        p = document.add_paragraph()
        r = p.add_run(text)

        for filename in os.listdir('./'):
            if filename.endswith('.jpg'):
                img = Image.open(filename)
                width = img.size[0]
                height = img.size[1]
                img.close()
                if width > height: # Album orientation 
                    r.add_picture(filename, width = Cm(12.0))
                else: # Portrait and quad
                    r.add_picture(filename, height = Cm(12.0))
                r.add_text(' ')
            else:
                continue

        os.chdir('../')

os.chdir(old_cwd)
document.save('Report.docx')
